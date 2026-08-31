#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""spec.py — 「规范查询」skill 唯一程序。

子命令(两态模型):
  维护态: index / ocr / status / remove / update-chars
  查询态: list / toc / clause / read / grep

设计要点:
  - 查询路径纯文件操作(读 toc.md / chapters / clauses.idx),零 Python 依赖;
  - OCR 只在索引时整本一次性执行,查询路径永不 OCR;
  - 质量检测用「常用字覆盖率 + 标点污染率」双信号识别乱码(两种乱码模式都覆盖);
  - 所有输出 UTF-8;未索引/缺缓存时明确报错并提示下一步,退出码非 0。
"""
import argparse
import json
import math
import os
import re
import subprocess
import sys
import tempfile
from contextlib import contextmanager
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

import fitz  # PyMuPDF

SKILL_DIR = Path(__file__).resolve().parent.parent
DEFAULT_CONFIG = SKILL_DIR / "config.json"

# ---------- 正则常量 ----------
ASCII_PUNCT = set('!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~')

# ---------- 支持的文件格式(库目录索引范围;.zip 等压缩包暂不索引) ----------
INDEXABLE_EXTS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".wps", ".et",
                  ".ofd", ".png", ".tif", ".tiff", ".txt", ".md"}
IMAGE_EXTS = {".png", ".tif", ".tiff"}      # 图片书: 整图/逐帧 tesseract,type=ocr
COM_EXTS = {".doc", ".xls", ".wps", ".et"}  # 老格式: COM 另存 docx/xlsx 再解析(串行)
VIRTUAL_EXTS = {".doc", ".docx", ".xls", ".xlsx", ".wps", ".et", ".txt", ".md"}  # 虚拟分页
# fmt(无点扩展名)与带点集合的成员判定;dispatch 处 fmt 均为无点形式
_IMAGE_FMTS = {e.lstrip(".") for e in IMAGE_EXTS}
_COM_FMTS = {e.lstrip(".") for e in COM_EXTS}
_VIRTUAL_FMTS = {e.lstrip(".") for e in VIRTUAL_EXTS}
SKIP_HINTS = {".zip": "zip 压缩包暂不索引(需解包后放入库目录)"}
VIRT_TARGET = 500        # 每虚拟页目标字符数
VIRT_MAX_PAGES = 400     # 虚拟页数硬顶(防御大 xlsx 等)

# COM 单例(仅在 cmd_index 的 COM 串行阶段使用,严禁进线程池)
_COM_CTX = None

# 条文号: 行首 X.Y 或 X.Y.Z,后跟空白/行尾/汉字(GB 两层、JTG 三层、同行粘连、全角空格都兼容)
CLAUSE_RE = re.compile(
    r'^[ \t　\xa0]*(\d{1,2})(?:[ \t　\xa0]*\.[ \t　\xa0]*(\d{1,2}))(?:[ \t　\xa0]*\.[ \t　\xa0]*(\d{1,3}))?'
    r'(?=[ \t　\xa0]*$|[ \t　\xa0]*[一-鿿])')
# 裸数字行(条下列项 "1"/"2"),不是条文
BARE_NUM_RE = re.compile(r'^[ \t　\xa0]*\d{1,2}[ \t　\xa0]*$')


def _spaced_variant_ok(s, m):
    """空格变体行(条文号部分含提取间距,如 "4. 2. 12")内容校验:
    幽灵条文(量词行/短节标题)与伪交叉引用行(引用连词开头)跳过,真条文保留。
    2 级(X.Y)在 GB 两层编号书里就是真条文号,不能一律丢弃——只有量词/连词/
    短标题(≤6 字无标点)形态判定为幽灵("1. 2 标题"/"3 . 5 米"/"4. 2 条规定");
    长正文("4. 2 桥梁养护的内容…")保留。3 级(X.Y.Z)只拒量词/连词,条文标题
    ("4. 2. 1 一般规定")保留。连词集合限高判别力(按/符合/依据常见于真条文
    正文开头,如 "符合下列规定时…",不拒)。"""
    rest = s[m.end():].lstrip(' \t　\xa0')
    # 量词/数字开头(米/倍/条/次/天/年;实测"个/元"词首误伤真条文已剔除)
    if not rest or rest[0].isdigit() or rest[0] in '米倍条次天年％%':
        return False
    # 引用连词开头(和/与/及/或/图/表/见/至/同)
    if rest[0] in '和与及或图表见至同':
        return False
    # 2 级短标题/无标点短行 = 节编号("1. 2 标题"/"4. 2 一般规定")
    if m.group(3) is None and not re.search(r'[，。；：、]', rest) \
            and len(re.sub(r'[\s　\xa0]+', '', rest)) <= 6:
        return False
    return True


def _clause_match_ok(s):
    """CLAUSE_RE 匹配 + 空格变体行过滤(幽灵条文/伪交叉引用)。返回 m 或 None。
    build_clauses / _clause_prefix_first_pages / _is_expl_page 共用。"""
    m = CLAUSE_RE.match(s)
    if m and re.search(r'\.[ \t　\xa0]+', s[:m.end()]) and not _spaced_variant_ok(s, m):
        return None
    return m
# 目录行解析见 _parse_toc_line(宽容版: 页码可有可无)
# 章标题: 数字+空白+汉字(同行式 "4 桥梁养护与维修")
CHAPTER_INLINE_RE = re.compile(r'^[ \t　\xa0]*(\d{1,2})[ \t　\xa0]+([一-鿿][^\n]*)')
# 章标题: 数字直接接汉字(粘连式 "7二、三级…")
CHAPTER_GLUE_RE = re.compile(r'^[ \t　\xa0]*(\d{1,2})([一-鿿][^\n]*)')
# 章标题: 独立数字行(下一非空行是中文标题)
CHAPTER_STANDALONE_RE = re.compile(r'^[ \t　\xa0]*(\d{1,2})[ \t　\xa0]*$')
APPENDIX_RE = re.compile(r'^[ \t　\xa0]*(附录[A-Z])[ \t　\xa0]*(.*)$')
TOC_MARKER_RE = re.compile(r'目\s*[次录]')
EXPL_MARKER_RE = re.compile(r'^\s*条\s*文\s*说\s*明')
EXPL_TITLE_RE = re.compile(r'^[ \t　\xa0]*条[ \t　\xa0]*文[ \t　\xa0]*说[ \t　\xa0]*明[ \t　\xa0]*$')

# 全角数字/变体点单源(文字版条文号 "１􀆰０􀆰１" = 1.0.1);FULLWIDTH_NORM 与
# OCR_NORM_TABLE 共用,防两表漂移。注意不含 '〇'——正文日期 "二〇二六年" 是
# 合法文字(整页翻译会破坏原文),仅 OCR_NORM_TABLE 用于条文号/目录行短上下文
_FULLWIDTH_MAP = {'１': '1', '２': '2', '３': '3', '４': '4',
                  '５': '5', '６': '6', '７': '7', '８': '8', '９': '9',
                  '０': '0', '．': '.', '􀆰': '.'}
OCR_NORM_TABLE = str.maketrans({'l': '1', '|': '1', '丨': '1', 'I': '1', '丿': '1',
                                'O': '0', 'o': '0', '〇': '0', 'D': '0',
                                's': '5', 'S': '5', '$': '', '^': '',
                                **_FULLWIDTH_MAP})
# 章节文件正文用(不含 OCR 混淆映射 l→1 等,不含 '〇'——见 _FULLWIDTH_MAP 注释):
# 全角数字/变体点 → 半角,使 grep 章节文件与 clause 直查一致
FULLWIDTH_NORM = str.maketrans(_FULLWIDTH_MAP)
ILLEGAL_FNAME = re.compile(r'[\\/:*?"<>|\r\n\t]')


def die(msg):
    sys.stderr.write(f"[spec.py] {msg}\n")
    sys.exit(1)


def is_cjk(ch):
    o = ord(ch)
    return 0x4E00 <= o <= 0x9FFF or 0x3400 <= o <= 0x4DBF or 0xF900 <= o <= 0xFAFF


def median(xs):
    if not xs:
        return None
    s = sorted(xs)
    n = len(s)
    return s[n // 2] if n % 2 else (s[n // 2 - 1] + s[n // 2]) / 2


# ---------- 配置 / 书架 ----------

def load_config(args):
    cfg_path = Path(args.config) if getattr(args, "config", None) else DEFAULT_CONFIG
    try:
        cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        die(f"config.json 不存在: {cfg_path}")
    cfg["_path"] = cfg_path
    return cfg


def atomic_write_text(path, text, encoding="utf-8"):
    """同目录临时文件 + replace，避免中断时留下截断的索引/书架文件。"""
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    tmp_name = None
    try:
        with tempfile.NamedTemporaryFile("w", encoding=encoding, dir=p.parent,
                                         prefix=f".{p.name}.", suffix=".tmp",
                                         delete=False) as f:
            tmp_name = f.name
            f.write(text)
            f.flush()
            os.fsync(f.fileno())
        os.replace(tmp_name, p)
    finally:
        if tmp_name and os.path.exists(tmp_name):
            os.unlink(tmp_name)


@contextmanager
def library_lock(data_dir):
    """同一索引库只允许一个维护命令写入，防止并发进程互相覆盖书架。"""
    root = Path(data_dir)
    root.mkdir(parents=True, exist_ok=True)
    lock = root / ".spec.lock"
    try:
        fd = os.open(lock, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
    except FileExistsError:
        die(f"索引库正在被另一维护命令使用: {lock}。确认无运行任务后删除该锁文件再重试")
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(f"pid={os.getpid()}\n")
        yield
    finally:
        try:
            lock.unlink()
        except FileNotFoundError:
            pass


def book_data_dir(cfg, b):
    """索引数据目录 = data_dir / 源文件相对父目录 / book_id(与 guifansrc 目录结构一致)。
    gonglu/xxx.pdf → library_data/gonglu/<book_id>/;根目录书 → library_data/<book_id>/"""
    return Path(cfg["data_dir"]) / Path(b.get("file") or "").parent / b["id"]


def _src_abs(meta):
    """源文件绝对路径: 新字段 source_abs 优先,回退旧字段 pdf_abs(兼容期)。"""
    return meta.get("source_abs") or meta.get("pdf_abs")


def _src_mtime(entry):
    """源文件 mtime: source_mtime 优先,回退 pdf_mtime(兼容期)。"""
    return entry.get("source_mtime") or entry.get("pdf_mtime")


def source_signature(path):
    """增量索引的源文件指纹；纳秒 mtime + 大小避免秒级时间戳漏检。"""
    st = Path(path).stat()
    return {"mtime_ns": st.st_mtime_ns, "size": st.st_size,
            "display_mtime": datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M:%S")}


def entry_matches_source(entry, sig):
    """兼容旧书架：旧记录仍按展示用 mtime 比较，新记录使用精确指纹。"""
    if "source_mtime_ns" in entry or "source_size" in entry:
        return (entry.get("source_mtime_ns") == sig["mtime_ns"]
                and entry.get("source_size") == sig["size"])
    return _src_mtime(entry) == sig["display_mtime"]


def load_shelf(data_dir):
    p = Path(data_dir) / "bookshelf.json"
    if not p.exists():
        return {"schema_version": 1, "books": []}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception as e:
        die(f"bookshelf.json 读取失败: {e}")


def save_shelf(data_dir, shelf):
    p = Path(data_dir) / "bookshelf.json"
    atomic_write_text(p, json.dumps(shelf, ensure_ascii=False, indent=2) + "\n")


def find_book(shelf, key):
    """按 id / 文件名 / 规范号 / 书名 匹配书架中的书。"""
    key_l = key.strip().lower()
    cands = []
    for b in shelf.get("books", []):
        if b["id"].lower() == key_l:
            return b
        if key_l in (b.get("file", "").lower(), norm_no(b.get("std_no", "")), b.get("title", "").lower()):
            cands.append(b)
    if len(cands) == 1:
        return cands[0]
    for b in shelf.get("books", []):
        if key_l in b.get("title", "").lower() or key_l in norm_no(b.get("std_no", "")):
            cands.append(b)
    uniq = {b["id"]: b for b in cands}
    if len(uniq) == 1:
        return list(uniq.values())[0]
    if len(uniq) > 1:
        die(f"「{key}」匹配到多本书: {', '.join(uniq)} — 请用更精确的 id/规范号")
    die(f"书架中找不到「{key}」。先 `python spec.py list` 看书架,新书先 `spec.py index <文件>`")


def norm_no(s):
    # \s 不含 U+200B 等零宽(Cf 类),PDF 提取/复制可能插入——一并去掉
    return re.sub(r'[\s​‌‍﻿]+', '', str(s or '')).lower()


def slug_id(s):
    s = norm_no(s)
    s = s.replace('.', '-')
    s = re.sub(r'[^a-z0-9\-一-鿿]+', '-', s).strip('-')
    s = re.sub(r'-+', '-', s)
    return s or 'book'


def parse_filename(name):
    """从文件名解析序号/书名/规范号: 1.公路桥涵养护规范(JTG 5120-2021).pdf"""
    stem = Path(name).stem.strip()
    m = re.match(r'^(\d+)[.、．]\s*', stem)
    seq = int(m.group(1)) if m else None
    stem = re.sub(r'^\d+[.、．]\s*', '', stem)
    stem = stem.replace('+', ' ')
    m = re.search(r'^(.*?)[（(]([^（）()]*)[）)]$', stem)
    if m:
        title = m.group(1).strip()
        std_no = m.group(2).strip()
    else:
        title = stem.strip()
        std_no = ''
    version = ''
    mm = re.search(r'(\d{4})', std_no)
    if mm:
        version = mm.group(1)
    return title, std_no, version, seq


def make_book_id(seq, title, std_no):
    """(兼容保留) 旧规则: 序号-名称-编号。新书 book_id 直接用源文件名(去 .pdf)。"""
    parts = [f"{seq:02d}" if seq else '', title.strip()]
    if std_no:
        parts.append(std_no)
    return '-'.join(p for p in parts if p)


def _rel_source(src_path, cfg):
    """源文件相对 library_dir 的路径(正斜杠),支持多层文件夹。"""
    lib = Path(cfg["library_dir"]).resolve()
    p = src_path.resolve()
    try:
        return p.relative_to(lib).as_posix()
    except ValueError:
        return src_path.name


def _is_zip(path):
    """魔数检测: 文件是否为 ZIP(PK 头;xlsx/docx 必须,doc/xls 若是则多为改名文件)。"""
    try:
        with open(path, "rb") as f:
            return f.read(2) == b"PK"
    except OSError:
        return False


def _short_id(base):
    """book_id 超长截断(60+sha1): 目录内附件文件名 = 完整标题,与父目录名叠加
    超 Windows 路径上限;截断后 id 不再等于文件名,但 find_book 按 file 字段仍可匹配。"""
    if len(base) <= 80:
        return base
    import hashlib
    return base[:60] + "_" + hashlib.sha1(base.encode("utf-8")).hexdigest()[:8]


def _path_book_id(rel, base, used_ids):
    """同名源文件的稳定 ID。父目录同名时仍由完整相对路径哈希保证唯一。"""
    import hashlib
    stem = f"{base}_{hashlib.sha1(rel.encode('utf-8')).hexdigest()[:10]}"
    candidate = _short_id(stem)
    # SHA1 前缀碰撞极罕见，但书架 ID 是主键，仍显式兜底。
    n = 1
    while candidate in used_ids:
        candidate = _short_id(f"{stem}_{n}")
        n += 1
    return candidate


def plan_book_ids(sources, cfg, shelf):
    """在进入线程池前预分配 book_id，避免同一批同名文件互相覆盖。"""
    existing_by_file = {b.get("file"): b for b in shelf.get("books", [])}
    used_ids = {b.get("id") for b in shelf.get("books", []) if b.get("id")}
    pending = []
    for src in sources:
        rel = _rel_source(Path(src), cfg)
        if rel not in existing_by_file:
            pending.append((rel, Path(rel).stem.rstrip(' .')))
    counts = Counter(base for _, base in pending)
    existing_stems = {Path(b.get("file") or "").stem.rstrip(' .')
                      for b in shelf.get("books", [])}
    plan = {rel: b["id"] for rel, b in existing_by_file.items() if b.get("id")}
    for rel, base in sorted(pending):
        # 独占文件保留易读的文件名 ID；只要出现同 stem(含后续新增)则使用路径哈希。
        if counts[base] == 1 and base not in existing_stems and base not in used_ids:
            book_id = base
        else:
            book_id = _path_book_id(rel, base, used_ids)
        plan[rel] = book_id
        used_ids.add(book_id)
    return plan


def scan_library(lib_dir):
    """收集库目录所有可索引文件 → (indexable: {rel: abs}, hints: [(name, 原因)])。
    index --all 与 status 共用,保证扩展名口径一致;zip/无扩展名等只提示不索引。"""
    lib = Path(lib_dir).resolve()
    indexable, hints = {}, []
    for p in lib.rglob("*"):
        try:
            if not p.is_file():
                continue
        except OSError:
            hints.append((p.name, "超长文件名无法读取,跳过"))
            continue
        ext = p.suffix.lower()
        try:
            rel = p.relative_to(lib).as_posix()
        except ValueError:
            rel = p.name
        if ext in INDEXABLE_EXTS:
            indexable[rel] = str(p)
        elif ext in SKIP_HINTS:
            hints.append((p.name, SKIP_HINTS[ext]))
        elif not ext:
            hints.append((p.name, "无扩展名,无法识别格式"))
    return indexable, hints


# ---------- 质量检测 ----------

def load_chars(path):
    p = Path(path)
    if not p.is_absolute():
        p = SKILL_DIR / path
    if not p.exists():
        die(f"常用字表缺失: {p}。先 `spec.py update-chars --from-pdfs <干净文字版PDF...>` 生成")
    chars = set()
    for line in p.read_text(encoding="utf-8").splitlines():
        for ch in line.strip():
            if is_cjk(ch):
                chars.add(ch)
    return chars


def page_stats(text, chars):
    n = len(text)
    cjk = [c for c in text if is_cjk(c)]
    cov = len([c for c in cjk if c in chars]) / len(cjk) if cjk else None
    nonws = [c for c in text if not c.isspace()]
    punct = len([c for c in nonws if c in ASCII_PUNCT]) / len(nonws) if nonws else None
    return {"n": n, "cov": cov, "punct": punct}


def probe_pdf(pdf_path, chars):
    """质量检测与分流: 返回 type 与证据。采样前 5 页 + 均匀 10 页,取中位数。"""
    doc = fitz.open(pdf_path)
    pages = doc.page_count
    sample = list(dict.fromkeys(
        [1, 2, 3, 4, 5] +
        [round(1 + i * (pages - 1) / 9) for i in range(10) if pages > 1] +
        [pages]))
    samples = [(p, page_stats(doc[p - 1].get_text(), chars))
               for p in sample if 1 <= p <= pages]
    # 扫描件判定用全部采样页(含封面/目录)
    empty_ratio = len([1 for _, s in samples if s["n"] < 20]) / len(samples) if samples else 1
    # 覆盖/污染只看正文区(p>=6): 封面/公告/目录页常有乱码或噪声,不参与
    body = [(p, s) for p, s in samples if p >= 6 and s["n"] >= 20]
    covs = [s["cov"] for _, s in body if s["cov"] is not None]
    puncts = [s["punct"] for _, s in body if s["punct"] is not None]
    cov_med = median(covs) if covs else None
    punct_med = median(puncts) if puncts else None
    # 劣化页: 有内容但覆盖率低或标点污染高
    bad = [s for _, s in body if
           s["cov"] is None or s["cov"] < 0.95 or (s["punct"] is not None and s["punct"] > 0.08)]
    bad_ratio = len(bad) / len(body) if body else 1
    # 条文号命中仲裁: 前 30% 页(至少 3 页;min 上界防 1-2 页管理文件越界)
    hits = 0
    for p in range(1, min(pages, max(3, int(pages * 0.3))) + 1):
        if hits > 200:
            break
        for ln in doc[p - 1].get_text().splitlines():
            # 与 build_clauses 同口径(空格变体幽灵行不计入仲裁)
            if _clause_match_ok(ln.translate(OCR_NORM_TABLE)):
                hits += 1
                if hits > 200:
                    break
    doc.close()
    # 扫描判定阈值: 0.9→0.7,抓 70%+ 无字页的扫描件(如 012 护栏评价,78.6% 无字误判 text)
    is_scan = empty_ratio > 0.7
    is_garbled = (cov_med is not None and cov_med < 0.95) or (
        punct_med is not None and punct_med > 0.08 and bad_ratio > 0.6)
    typ = "ocr" if (is_scan or is_garbled) else "text"
    rule = []
    if is_scan:
        rule.append("scan")
    if is_garbled:
        rule.append("garbled")
    if not rule:
        rule.append("clean")
    # 仲裁: 疑似乱码但条文号命中充足且覆盖率尚可 → 救回 text
    if typ == "ocr" and hits > 50 and (cov_med or 0) >= 0.95:
        typ, rule = "text", rule + ["rescued-by-clause-hits"]
    return {"type": typ, "pages": pages, "rule": rule,
            "cov_median": round(cov_med, 4) if cov_med is not None else None,
            "punct_median": round(punct_med, 4) if punct_med is not None else None,
            "empty_ratio": round(empty_ratio, 3),
            "bad_ratio": round(bad_ratio, 3),
            "clause_hits": hits}


# ---------- 文本提取 / OCR ----------

def probe_text(texts, chars):
    """非 PDF 原生文本书的质量检测: 不做 OCR 分流(Office/OFD 提取即权威),
    逐页统计作证据记录;全空文档由上层挂 low_confidence 提示。"""
    stats = [page_stats(t, chars) for t in texts.values() if t]
    body = [s for s in stats if s["n"] >= 20]
    covs = [s["cov"] for s in body if s["cov"] is not None]
    puncts = [s["punct"] for s in body if s["punct"] is not None]
    return {"type": "text", "pages": len(texts), "rule": ["native-text"],
            "cov_median": round(median(covs), 4) if covs else None,
            "punct_median": round(median(puncts), 4) if puncts else None,
            "empty_ratio": round(1 - len(body) / len(texts), 3) if texts else 1,
            "bad_ratio": 0, "clause_hits": 0}


def virtual_paginate(lines):
    """非 PDF 文本书虚拟分页: 按段落累计 ~VIRT_TARGET 字符切页,段落不撕裂。
    页数自适应(大文档 target 上调)+ VIRT_MAX_PAGES 硬顶,防章节文件爆炸。"""
    if not lines:
        return [""]
    total = sum(len(ln) for ln in lines)
    target = max(VIRT_TARGET, math.ceil(total / 300))
    pages, cur, n = [], [], 0
    for ln in lines:
        if cur and n + len(ln) > target:
            pages.append("\n".join(cur))
            cur, n = [], 0
        cur.append(ln)
        n += len(ln) + 1
    if cur:
        pages.append("\n".join(cur))
    if len(pages) > VIRT_MAX_PAGES:
        pages = pages[:VIRT_MAX_PAGES - 1] + ["\n".join(pages[VIRT_MAX_PAGES - 1:])]
    return pages


def extract_docx(path):
    """docx → 文本行: 段落+表格按文档体顺序保序,页眉/页脚(红头单位名)前置。"""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    doc = Document(path)
    rows = []

    def _collect_header_footer():
        hf = []
        for sect in doc.sections:
            for p in list(sect.header.paragraphs) + list(sect.footer.paragraphs):
                t = p.text.strip()
                if (t and t not in hf and len(t) <= 40
                        and not re.fullmatch(r'[\s\d\-—第页共]*', t)):
                    hf.append(t)
        return hf

    rows.extend(_collect_header_footer())

    def _cell_text(cell):
        parts = [x.text.strip() for x in cell.paragraphs if x.text.strip()]
        return " ".join(parts)

    def _iter_blocks(parent):
        for child in parent.element.body.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                yield Table(child, parent)

    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            t = block.text.strip()
            if t:
                rows.append(t)
        else:
            for r in block.rows:
                cells = [_cell_text(c) for c in r.cells]
                cells = [c for c in cells if c]
                if cells:
                    rows.append("\t".join(cells))
    return rows


def extract_xlsx(path):
    """xlsx → 文本行: 每工作表一行【工作表 i:名称】标记 + 逐行 tab 连接。"""
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    rows = []
    for i, ws in enumerate(wb.worksheets, 1):
        rows.append(f"【工作表 {i}:{ws.title}】")
        for r in ws.iter_rows(values_only=True):
            cells = [str(v).strip() for v in r if v is not None and str(v).strip()]
            if cells:
                rows.append("\t".join(cells))
    wb.close()
    return rows


def extract_ofd(path):
    """ofd → 页文本列表: 每物理页 1 项(Content.xml 的 TextCode 文本层)。
    纯 Python(zipfile + ElementTree),不碰 COM。"""
    import zipfile
    import xml.etree.ElementTree as ET
    OFD_TAG = "{http://www.ofdspec.org/2016}"
    pages = []
    with zipfile.ZipFile(path) as z:
        page_files = sorted(
            (n for n in z.namelist()
             if re.search(r'/Page_\d+/Content\.xml$', n)),
            key=lambda n: int(re.search(r'/Page_(\d+)/', n).group(1)))
        for pn in page_files:
            root = ET.fromstring(z.read(pn))
            texts = []
            for obj in root.iter(OFD_TAG + "TextObject"):
                codes = [tc.text or "" for tc in obj.iter(OFD_TAG + "TextCode")]
                t = "".join(codes).strip()
                if t:
                    texts.append(t)
            pages.append("\n".join(texts))
    if not pages:
        pages = [""]
    return pages


def extract_txt(path):
    """txt/md → 文本行(编码容错)。"""
    return [ln.strip() for ln in path.read_text(encoding="utf-8", errors="replace")
            .splitlines() if ln.strip()]


def image_page_count(path):
    """图片书页数: PNG=1,TIF/TIFF=帧数。"""
    if path.suffix.lower() == ".png":
        return 1
    from PIL import Image
    with Image.open(path) as im:
        return max(1, getattr(im, "n_frames", 1))


def _image_frame_bytes(path, frame):
    """图片书第 frame 帧的 PNG 字节(PNG=整文件;TIF 逐帧经 Pillow 转 PNG)。"""
    if path.suffix.lower() == ".png":
        return path.read_bytes()
    from PIL import Image
    import io
    with Image.open(path) as im:
        im.seek(frame - 1)
        buf = io.BytesIO()
        im.convert("RGB").save(buf, format="PNG")
        return buf.getvalue()


def _write_extracted(book_dir, texts):
    """非 PDF 书文本落盘 extracted/NNN.txt(read 查询读这里;OCR 图片书另有 ocr/)。"""
    d = Path(book_dir) / "extracted"
    d.mkdir(parents=True, exist_ok=True)
    for p, t in texts.items():
        atomic_write_text(d / f"{p:03d}.txt", t)


def extract_text_pages(doc, book_dir, page_count):
    """text 书: 逐页提取(仅内存,不落盘 pages/),返回 {页: 文本}"""
    texts = {}
    for i in range(page_count):
        texts[i + 1] = doc[i].get_text()
    return texts


def run_tesseract(png_bytes, cfg):
    cmd = [cfg.get("tesseract_cmd", "tesseract"), "stdin", "stdout",
           "-l", cfg.get("ocr_lang", "chi_sim"), "--psm", str(cfg.get("ocr_psm", 3))]
    td = cfg.get("ocr_tessdata_dir")
    if td:
        p = Path(td)
        if not p.is_absolute():
            p = SKILL_DIR / p
        cmd += ["--tessdata-dir", str(p)]
    r = subprocess.run(cmd, input=png_bytes, capture_output=True, timeout=600)
    return r.returncode, r.stdout.decode("utf-8", errors="replace")


def run_ocr_pages(src, book_dir, cfg, start, end, force, fmt):
    """整本批量 OCR: PDF 300dpi 渲染 / 图片书整图逐帧 → tesseract → ocr/NNN.txt。
    断点续跑,写缓存跳过,--force 重跑。"""
    ocr_dir = Path(book_dir) / "ocr"
    ocr_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(src) if fmt == "pdf" else None
    dpi = int(cfg.get("ocr_dpi", 300))
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    done, failed = 0, []
    total = end - start + 1
    for p in range(start, end + 1):
        out = ocr_dir / f"{p:03d}.txt"
        if out.exists() and not force:
            continue
        png = (doc[p - 1].get_pixmap(matrix=mat).tobytes("png") if fmt == "pdf"
               else _image_frame_bytes(src, p))
        ok = False
        for attempt in (1, 2):
            try:
                rc, txt = run_tesseract(png, cfg)
                if rc == 0:
                    atomic_write_text(out, txt)
                    ok = True
                    break
            except Exception:
                pass
        if not ok:
            failed.append(p)
        done += 1
        if done % 10 == 0:
            print(f"  ocr {done}/{total} ({100 * done // total}%)", flush=True)
    if doc:
        doc.close()
    return failed


def _ocr_book(src, book_dir, cfg, pages, label, fmt, force=False):
    """整本 OCR + 读回文本(断点续跑:缓存跳过,--force 不重 OCR)。
    返回 (failed_pages, texts)。PDF 与图片书共用。"""
    failed = run_ocr_pages(src, book_dir, cfg, 1, pages, force, fmt)
    if failed:
        # label 带 [book_id] 前缀(--jobs 并行时失败日志可归属具体书)
        print(f"[{label}] OCR 失败 {len(failed)} 页: {failed[:10]}...", flush=True)
    return failed, load_book_texts(book_dir, pages, "ocr")


class ComContext:
    """COM 单例: 每 ProgID 一个 Application 多文档复用;连续失败重建;退出 QuitAll。
    COM 是 apartment 线程模型,本类只在主线程/COM 串行阶段使用,严禁进线程池。"""

    _PROGIDS = {"word": ("Word.Application", "docx"),
                "excel": ("Excel.Application", "xlsx"),
                "wps": ("kwps.Application", "docx"),
                "et": ("ket.Application", "xlsx")}
    _FMT_DOCX, _FMT_XLSX = 12, 51  # wdFormatXMLDocument / xlOpenXMLWorkbook

    def __init__(self):
        self._apps = {}
        self._fails = Counter()
        self._tmpdir = None
        self._pre_pids = self._snapshot_pids()

    def _get_app(self, kind):
        app = self._apps.get(kind)
        if app is not None:
            return app
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()
        app = win32com.client.DispatchEx(self._PROGIDS[kind][0])
        for attr in ("Visible",):
            try:
                setattr(app, attr, False)
            except Exception:
                pass
        try:
            app.DisplayAlerts = 0
        except Exception:
            pass
        if kind in ("word", "wps"):  # 禁宏,防转换弹窗
            try:
                app.AutomationSecurity = 3
            except Exception:
                pass
        self._apps[kind] = app
        return app

    def convert(self, src, kind):
        """打开 → 另存临时 .docx/.xlsx → 返回临时文件路径。失败抛异常。"""
        app = self._get_app(kind)
        if self._tmpdir is None:
            import tempfile
            self._tmpdir = Path(tempfile.mkdtemp(prefix="spec_com_"))
        import hashlib
        name = src.stem[:60] + "_" + hashlib.sha1(str(src).encode("utf-8")).hexdigest()[:8]
        dest = self._tmpdir / (name + "." + self._PROGIDS[kind][1])
        try:
            if kind in ("word", "wps"):
                doc = app.Documents.Open(str(src), ReadOnly=True, AddToRecentFiles=False)
                try:
                    doc.SaveAs2(str(dest), FileFormat=self._FMT_DOCX)
                except Exception:
                    doc.SaveAs(str(dest), FileFormat=self._FMT_DOCX)
                doc.Close(False)
            else:
                wb = app.Workbooks.Open(str(src), ReadOnly=True, UpdateLinks=0)
                try:
                    wb.SaveAs(str(dest), FileFormat=self._FMT_XLSX)
                except Exception:
                    wb.SaveAs(str(dest), FileFormat=self._FMT_XLSX)
                wb.Close(False)
            self._fails[kind] = 0
            return dest
        except Exception as e:
            self._fails[kind] += 1
            if self._fails[kind] >= 2:  # 连续失败 → 重建 Application(防状态污染连带后续)
                self._quit(kind)
            raise
        finally:
            pass

    @staticmethod
    def _snapshot_pids():
        """当前 wps/pet 等套件进程 PID(Quit 后差分清理,避免误杀用户已开的 WPS)。"""
        pids = set()
        for name in ("wps.exe", "pet.exe", "et.exe", "wpp.exe"):
            try:
                r = subprocess.run(["tasklist", "/FI", f"IMAGENAME eq {name}",
                                    "/FO", "CSV", "/NH"], capture_output=True,
                                   text=True, timeout=30)
                for ln in r.stdout.splitlines():
                    m = re.match(r'^"([^"]+)","(\d+)"', ln.strip())
                    if m:
                        pids.add(int(m.group(2)))
            except Exception:
                pass
        return pids

    def _kill_leftover(self):
        """Quit 后 WPS 套件常残留子进程(pet 等),清掉本次新建的(差分)。"""
        for pid in sorted(self._snapshot_pids() - self._pre_pids):
            try:
                subprocess.run(["taskkill", "/PID", str(pid), "/F"],
                               capture_output=True, timeout=30)
            except Exception:
                pass

    def _quit(self, kind):
        app = self._apps.pop(kind, None)
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass

    def quit_all(self):
        for kind in list(self._apps):
            self._quit(kind)
        self._kill_leftover()
        if self._tmpdir is not None:
            import shutil
            shutil.rmtree(self._tmpdir, ignore_errors=True)
            self._tmpdir = None


def extract_word_com(src, cfg, com):
    """doc/wps → COM 另存临时 docx → extract_docx(单一解析路径,表格结构保留)。
    MS Office 拒绝打开时跨应用用 WPS 兜底(部分损坏/老格式文档 Office 拒开而 WPS 能开)。"""
    kind = "word" if src.suffix.lower() == ".doc" else "wps"
    try:
        tmp = com.convert(src, kind)
    except Exception:
        if kind != "wps":
            tmp = com.convert(src, "wps")
        else:
            raise
    return extract_docx(tmp)


def extract_excel_com(src, cfg, com):
    """xls/et → COM 另存临时 xlsx → extract_xlsx。Excel 拒绝时 WPS 表格(ket)兜底。"""
    kind = "excel" if src.suffix.lower() == ".xls" else "et"
    try:
        tmp = com.convert(src, kind)
    except Exception:
        if kind != "et":
            tmp = com.convert(src, "et")
        else:
            raise
    return extract_xlsx(tmp)


def load_book_texts(book_dir, page_count, kind):
    """从 pages/ 或 ocr/ 读回全书文本 {页: 文本}。"""
    sub = "ocr" if kind == "ocr" else "pages"
    base = Path(book_dir) / sub
    texts = {}
    for p in range(1, page_count + 1):
        f = base / f"{p:03d}.txt"
        if f.exists():
            texts[p] = f.read_text(encoding="utf-8")
    return texts


# ---------- 目录(TOC)三路径 ----------

def _resolve_dest_page(doc, dest):
    """书签页码为 -1 时,尝试解析 xref 目标。失败返回 None。"""
    try:
        if isinstance(dest, dict):
            to = dest.get("to")
            if to:
                obj = doc.xref_object(to)
                m = re.search(r'/D\s*\[\s*(\d+)\s+\d+\s+R', obj)
                if m:
                    xref = int(m.group(1))
                    for i in range(doc.page_count):
                        if doc.page_xref(i) == xref:
                            return i + 1
            d = dest.get("d")
            if isinstance(d, (list, tuple)) and d and isinstance(d[0], int):
                for i in range(doc.page_count):
                    if doc.page_xref(i) == d[0]:
                        return i + 1
    except Exception:
        pass
    return None


def toc_from_bookmarks(doc):
    """路径 1: 书签 level==1。任一页 -1 且解析失败 → 判定不可用。"""
    try:
        toc = doc.get_toc(simple=False)
    except Exception:
        return None
    if not toc:
        return None
    entries = []
    for lvl, title, page, dest in toc:
        if lvl != 1:
            continue
        if page < 1:
            page = _resolve_dest_page(doc, dest)
            if page is None:
                return None
        entries.append((title.strip(), page))
    if not entries:
        return None
    # 书签标题质量校验: 正常书签标题 CJK 占比高;乱码/占位书签(如 "WQ.pdf"/"ZW")
    # 视为损坏弃用,回退目录页解析/正文扫描
    total_len = sum(len(t) for t, _ in entries)
    cjk_n = sum(1 for t, _ in entries for c in t if is_cjk(c))
    if len(entries) < 3 or total_len == 0 or cjk_n / total_len < 0.3:
        return None
    return entries


def _bm_entries(entries):
    """书签条目 (title, page) → (label, title, page) 统一格式。"""
    out = []
    for title, pg in entries:
        m = re.match(r'^[ \t　\xa0]*(附录[A-Z]|[0-9]{1,2}(?:\.[0-9]{1,2}){0,2})[ \t　\xa0]*(.*)$', title)
        if m:
            out.append((m.group(1), m.group(2).strip() or m.group(1), pg))
        else:
            out.append(('', title, pg))
    return out


# 正文页信号: 条文号后跟长正文(≥12 汉字)——目录条目标题短,正文条文长
# 数字/点间兼容空格(与 CLAUSE_RE 一致,142/167 的 "1. 0. 1" 提取间距)
BODY_CLAUSE_LINE_RE = re.compile(
    r'^[ \t　\xa0]*\d{1,2}(?:[ \t　\xa0]*\.[ \t　\xa0]*\d{1,2}){1,2}[ \t　\xa0]+[一-鿿][^。\n]{11,}')


def _is_body_page(t):
    # 排除目录行: 长标题+点线+"页码"("3.1 桥梁抗震设防分类……9")是目录特征,
    # 会被 BODY_CLAUSE_LINE_RE 误判为正文条文行,导致目录页没被识别(015/065)
    # 目录点线是行尾形态("3.1 桥梁抗震设防分类……9");正文省略号 "……" 在句中,
    # 不能排除(否则含省略号的正文页失去正文信号,目录区延伸吞页)
    return len([ln for ln in t.splitlines()
                if BODY_CLAUSE_LINE_RE.match(ln)
                and not re.search(r'[.．·…]{2,}\s*\d{0,3}\s*$', ln)]) >= 3


def _parse_toc_line(ln, page_count):
    """宽容解析目录行: 编号 标题 [页码],页码可有可无("- 1 -"、点线+数字、纯点线)。"""
    ln = ln.translate(OCR_NORM_TABLE)  # OCR 数字混淆归一化(l→1 等)
    m = re.match(r'^[ \t　\xa0]*(附录[A-Z]|[0-9]{1,2}(?:\.[0-9]{1,2}){0,2})', ln)
    if not m:
        return None
    label = m.group(1)
    rest = ln[m.end():]
    stripped = rest.rstrip(' .…·-　\xa0\t')
    pg = None
    mm = re.search(r'(\d{1,3})\s*$', stripped)
    if mm:
        v = int(mm.group(1))
        if 1 <= v <= page_count:
            pg = v
            stripped = stripped[:mm.start()]
    title = stripped.strip(' .…·-　\xa0\t')
    if len(title) < 2 or not re.match(r'^[一-鿿]', title):
        return None
    # 前言等散文页的伪目录行(如"4 路基养护、5 路面养护…"起草分工句): 无页码且含数字/过长 → 拒
    if pg is None and (re.search(r'\d', title) or len(title) > 30):
        return None
    return label, title, pg


def _locate_heading(texts, label, title, lo, hi):
    """在正文中按 章号+标题 定位页(兼容独立行/同行/粘连三种排版)。"""
    def norm(s):
        return re.sub(r'[\s　\xa0]+', '', s)

    label_c, title_c = norm(label), norm(title)
    for p in range(max(1, lo), min(hi, 10 ** 9) + 1):
        lines = [ln.strip() for ln in texts.get(p, "").splitlines()]
        for i, ln in enumerate(lines):
            if not ln:
                continue
            if ln == label_c and i + 1 < len(lines) and norm(lines[i + 1]) == title_c:
                return p
            if norm(ln).startswith(label_c + title_c):
                return p
        for ln in lines:
            if norm(ln) == title_c:
                return p
    return None


def find_toc_zone(texts, page_count):
    """定位目录页区,返回 (start, end): 目录页 = 有"目次/目录"标记 且
    (≥3 条目录行 或 ≥5 条条文号行——OCR 书目录页点线被识别坏时的兜底)。
    end 连续延伸(带标记的后续页),遇正文页(≥3 条长条文行)或超 5 页即停。"""
    start = end = 0
    for p in range(1, min(16, page_count) + 1):
        t = texts.get(p, "")
        if _is_body_page(t):
            if start:
                break
            continue
        toc_hits = sum(1 for ln in t.splitlines() if _parse_toc_line(ln, page_count))
        # 注意: findall 第二参是 pos 非 flags,且 $ 前瞻在无 MULTILINE 时只匹配串尾
        # → 按行 match 才算数(曾导致 clause_hits 恒为 0,目录页识别失效)
        clause_hits = sum(1 for ln in t.splitlines()
                          if _clause_match_ok(ln.translate(OCR_NORM_TABLE)))
        if not start:
            # 起始页必须带"目次/目录"标记(前言/封面页不算)
            if TOC_MARKER_RE.search(t) and (toc_hits >= 3 or clause_hits >= 5):
                start = end = p
            continue
        if p - start >= 5:
            break
        # 延续页(目录续页常无标记): 目录行密度,或 非正文页且条文号行多
        # (编号独立行排版的目录,如 "9.1" 行+标题行);正文页(长条文行)不算
        if toc_hits >= 2 or (not _is_body_page(t) and clause_hits >= 5):
            end = p
        else:
            break
    return start, end


def toc_from_tocpages(texts, page_count):
    """路径 2: 解析目录页(页码残缺/缺失的书,用章标题回正文定位补页)。"""
    zone_start, zone_end = find_toc_zone(texts, page_count)
    if not zone_start:
        return None
    found = []
    for p in range(zone_start, zone_end + 1):
        for ln in texts.get(p, "").splitlines():
            r = _parse_toc_line(ln, page_count)
            if r:
                found.append(r + (p,))
    chapters = [(l, t, pg, p) for l, t, pg, p in found if '.' not in l]
    seen, merged = set(), []
    for c in chapters:
        if c[0] not in seen:
            seen.add(c[0])
            merged.append(c)
    if len(merged) < 3:
        return None
    # 页码偏移校准: 用第一个带页码的条目(标题回正文定位)
    offset = None
    for l, t, pg, _ in merged:
        if pg:
            hit = _locate_heading(texts, l, t, zone_end + 1, min(zone_end + 60, page_count))
            if hit:
                offset = hit - pg
                break
    # 组装: 有页码的按偏移换算;缺页码的按顺序从上一章之后定位
    entries, last_pdf = [], zone_end
    for l, t, pg, _ in merged:
        pdf = pg + offset if (pg and offset is not None) else None
        if pdf is None or not (1 <= pdf <= page_count):
            pdf = _locate_heading(texts, l, t, last_pdf + 1, page_count)
        if pdf:
            entries.append((l, t, pdf))
            last_pdf = pdf
    if len(entries) < 3:
        return None
    return entries, zone_end, offset


def _clause_prefix_first_pages(texts, zone, pages):
    """条文号聚类: 第 N 章条文(N.x)首次出现的页 → {N: page}(跳过目录区)。"""
    first = {}
    for p in range(zone + 1, pages + 1):
        t = texts.get(p, "")
        if not t:
            continue
        for ln in t.splitlines():
            m = _clause_match_ok(ln.translate(OCR_NORM_TABLE))
            if m:
                n = int(m.group(1))
                if n not in first:
                    first[n] = p
    return first


def _heading_on_page(texts, page, label):
    """在该页找 "N 标题" 行,返回标题(OCR 书标题可能被误读,仅作参考)。"""
    for ln in texts.get(page, "").splitlines():
        s = ln.strip()
        s_n = ln.translate(OCR_NORM_TABLE).strip()
        m = re.match(r'^' + re.escape(label) + r'[ \t　\xa0]+([一-鿿][^\n]*)', s_n)
        if m:
            return s[len(label):].strip(' \t　\xa0.…·-')[:30]
    return None


def toc_from_bodyscan(texts, page_count, start_page):
    """路径 3: 正文扫描章标题(目录页乱码/页码残缺时兜底)。"""
    entries, expl_start = [], None
    for p in range(start_page, page_count + 1):
        lines = texts.get(p, "").splitlines()
        if expl_start is None and _is_expl_page(texts.get(p, ""), p, page_count):
            expl_start = p
        for i, ln in enumerate(lines):
            s = ln.strip()
            s_n = ln.translate(OCR_NORM_TABLE).strip()  # OCR 数字混淆归一化
            if not s or len(s) > 40:
                continue
            if CLAUSE_RE.match(s_n) or BARE_NUM_RE.match(s_n):
                continue
            m = CHAPTER_STANDALONE_RE.match(s_n)
            if m and len(s) <= 4:
                for j in range(i + 1, min(i + 4, len(lines))):
                    nxt = lines[j].strip()
                    if not nxt:
                        continue
                    if re.match(r'^[一-鿿]', nxt) and 2 <= len(nxt) <= 40:
                        entries.append((m.group(1), nxt, p))
                    break
                continue
            # 注意用 s(未归一化): OCR_NORM_TABLE 的 D→0 会把"附录D"变"附录0",[A-Z] 失配
            m2 = APPENDIX_RE.match(s)
            if m2:
                title2 = s[2:].strip() or '附录'
                # 过滤正文句被误当附录标题("附录A 的要求。"/"B)。")
                if title2 != '附录' and (
                        '。' in title2 or len(re.findall(r'[一-鿿]', title2)) < 2):
                    continue
                entries.append((m2.group(1), title2, p))
                continue
            m3 = CHAPTER_INLINE_RE.match(s_n) or CHAPTER_GLUE_RE.match(s_n)
            if m3:
                title3 = m3.group(2).strip()
                # 过滤引言/前言编号列表假章(标题长或含句号是正文句)
                # 过滤 "第N章" 占位伪章与单字伪章(表格行被误当章标题)
                if (2 <= len(re.sub(r'[\s　\xa0]+', '', title3)) <= 15 and '。' not in title3
                        and not re.match(r'^第\d{1,2}章$', title3)):
                    entries.append((m3.group(1), title3, p))
    # 同页 ≥4 个纯数字编号的"章" → 数字列表噪声(条文说明区引用编号行);
    # 附录(A/B/C/D)等非数字 label 不受此限(真附录簇可同页多个)
    per_page = Counter(e[2] for e in entries if e[0].isdigit())
    entries = [e for e in entries
               if not (e[0].isdigit() and per_page[e[2]] >= 4)]
    # 去重 + 编号单调校验(每号取首次,页码须递增)
    seen, dedup = set(), []
    for label, title, pg in entries:
        if label in seen:
            continue
        seen.add(label)
        dedup.append((label, title, pg))
    dedup.sort(key=lambda x: (x[2], int(x[0]) if x[0].isdigit() else 99))
    mono = []
    for e in dedup:
        # 允许同页多个章标题(短章常挤在同一页)
        if not mono or e[2] >= mono[-1][2]:
            mono.append(e)
    return (mono, expl_start) if 5 <= len(mono) <= 40 else (None, expl_start)


def _is_expl_page(text, page, page_count):
    """真正的条文说明起始页: 位于书后半部 + 页首 8 行内整行「条文说明」
    + 同页重新从第 1 章条文开始编号。排除目录页噪声与"条文说明随条文
    交错排版"的书(那种书条文说明内嵌,不需要独立章)。"""
    if page <= 0.5 * page_count:
        return False
    lines = [ln for ln in text.splitlines() if ln.strip()][:8]
    if not any(EXPL_TITLE_RE.match(ln) for ln in lines):
        return False
    for ln in text.splitlines():
        s_n = ln.translate(OCR_NORM_TABLE)
        m = CLAUSE_RE.match(s_n)
        if m and m.group(1) == '1':
            # 空格变体行仅排除量词/数字形态;2 级 "1. 1 条文" 是条文说明区
            # 重编号特征,不能按节标题过滤(否则 expl 独立章丢失)
            rest = s_n[m.end():].lstrip(' \t　\xa0')
            if rest and (rest[0].isdigit() or rest[0] in '米倍条次天年％%'):
                continue
            return True
    return False


def detect_expl_start(texts, page_count, zone_end):
    for p in range(zone_end + 1, page_count + 1):
        if _is_expl_page(texts.get(p, ""), p, page_count):
            return p
    return None


# ---------- 章节 / 条文 ----------

def build_chapters(entries, expl_start, page_count):
    """entries: [(label, title, pdf_page)] → chapters 列表(条文说明独立末章)。"""
    chapters = []
    for label, title, pg in entries:
        if '条文说明' in title:
            continue  # 后面统一追加
        if expl_start and pg >= expl_start:
            continue  # 条文说明部分自带的分章(书签中与正文同号),并入条文说明章
        chapters.append({"label": label, "title": title, "start": pg, "is_expl": False})
    chapters.sort(key=lambda c: c["start"])
    if expl_start:
        chapters.append({"label": "条文说明", "title": "条文说明",
                         "start": expl_start, "is_expl": True})
    # 保证页码递增(条文说明在最末,若其页码被前面的章盖过则后移)
    fixed = []
    prev = 0
    for c in chapters:
        # 允许同页多章(短章挤在同一页),只保证不倒退
        if c["start"] < prev:
            c["start"] = prev
        prev = c["start"]
        fixed.append(c)
    return fixed


def slug_title(label, title):
    if label == '条文说明' or '条文说明' in title:
        return 'tiaowenshuoming'
    s = re.sub(r'[^0-9a-zA-Z一-鿿]+', '', title or '')
    return s[:20] or 'chapter'


def clause_range(chapters, i, idx_rows):
    """章 i 的条文号范围(取该页区间的首尾条文号)。"""
    start = chapters[i]["start"]
    end = chapters[i + 1]["start"] - 1 if i + 1 < len(chapters) else 10 ** 9
    nums = []
    for no, page, *_ in idx_rows:
        if start <= int(page) <= end:
            nums.append(no)
    if not nums:
        return '-'
    return f"{nums[0]}-{nums[-1]}"


def write_toc_md(book_dir, meta, chapters):
    offset = meta.get("offset")
    virtual = meta.get("virtual")
    col = "页" if virtual else "PDF 页"
    lines = [f"# {meta.get('title')}({meta.get('std_no')}) 目录", ""]
    if offset:
        lines.append(f"> 页码说明: PDF 页 = 正文页 + {offset}(1-based);chapters 章文件内用【第 N 页】标 PDF 页")
    if virtual:
        lines.append("> 页码为虚拟页(按 ~500 字符切分,非原文档页码),read 页码同义")
    lines.append("")
    lines.append(f"| 序号 | 章 | 标题 | {col} | 正文页 | 章文件 |")
    lines.append("|---|---|---|---|---|---|")
    for i, c in enumerate(chapters, 1):
        fname = f"chapters/ch{i:02d}-{slug_title(c['label'], c['title'])}.md"
        printed = c["start"] - offset if offset else ''
        lines.append(f"| {i} | {c['label']} | {c['title']} | {c['start']} | {printed} | {fname} |")
    atomic_write_text(Path(book_dir) / "toc.md", "\n".join(lines) + "\n")


def write_chapters(book_dir, texts, meta, chapters, idx_rows):
    ch_dir = Path(book_dir) / "chapters"
    ch_dir.mkdir(parents=True, exist_ok=True)
    # 清掉旧章文件(重建时章数/边界可能变化;兼容旧 .txt 后缀,避免残留误导查询)
    for f in list(ch_dir.glob("*.txt")) + list(ch_dir.glob("*.md")):
        f.unlink()
    # 方案 A: 文字书不再保留按页缓存,清掉旧 pages/(OCR 书的 ocr/ 保留)
    pages_dir = Path(book_dir) / "pages"
    if pages_dir.exists():
        import shutil
        shutil.rmtree(pages_dir)
    written = []
    for i, c in enumerate(chapters, 1):
        # 1-based i: 有下一章(i < len)时 end = 下一章起点-1;否则到书末。
        # 旧代码 i+1 < len 把倒数第二章当最后一章,其文件吞入最后一章全部内容
        end = chapters[i]["start"] - 1 if i < len(chapters) else meta["pages"]
        end = max(end, c["start"])
        fname = f"ch{i:02d}-{slug_title(c['label'], c['title'])}.md"
        offset = meta.get("offset")
        header = [f"# 章 {c['label']}:{c['title']}",
                  f"# 规范:{meta.get('std_no')} ({meta.get('id')}) | 页码:PDF {c['start']}-{end}"
                  + (f" | 正文页 {c['start'] - offset}-{end - offset}" if offset else ""),
                  f"# 条文范围:{clause_range(chapters, i - 1, idx_rows)}",
                  f"# 类型:{'条文说明' if c['is_expl'] else '正文'}",
                  f"# 生成:spec.py index {datetime.now().strftime('%Y-%m-%d')} | 页内标记【第 N 页】用于回源",
                  ""]
        parts = []
        for p in range(c["start"], end + 1):
            t = texts.get(p)
            if t is None:
                continue
            if offset:
                parts.append(f"【第 {p} 页|正文页 {p - offset}】")
            else:
                parts.append(f"【第 {p} 页】")
            # 全角数字/变体点归一化(１􀆰０􀆰１→1.0.1): 章节文件与 grep 检索一致,
            # 否则 clause 直查命中而 grep 二次确认搜不到
            parts.append(t.strip().translate(FULLWIDTH_NORM))
            parts.append("")
        atomic_write_text(ch_dir / fname, "\n".join(header) + "\n" + "\n".join(parts))
        written.append(fname)
    return written


def build_clauses(texts, meta, offset):
    """条文号索引: 跳过目录区,条文说明区标 expl。"""
    rows = []
    zone = meta.get("toc_zone_end", 0)
    expl = meta.get("expl_start")
    for p in range(1, meta["pages"] + 1):
        if p <= zone:
            continue
        t = texts.get(p)
        if not t:
            continue
        page_rows = []
        lines = t.splitlines()
        for i, ln in enumerate(lines):
            s = ln.strip()
            if BARE_NUM_RE.match(s):
                continue
            s_norm = ln.translate(OCR_NORM_TABLE)
            m = _clause_match_ok(s_norm)
            if not m:
                # 2025 新版排版: 条文号拆成 "4." 行 + 下一行 "5"(= 4.5),跨行组合(低置信)
                m_solo = re.match(
                    r'^[ \t　\xa0]*(\d{1,2})\.[ \t　\xa0]*$', s_norm)
                if m_solo:
                    for j in range(i + 1, min(i + 3, len(lines))):
                        m2 = re.match(
                            r'^[ \t　\xa0]*(\d{1,3})(?=[ \t　\xa0]|$|[一-鿿])',
                            lines[j].translate(OCR_NORM_TABLE))
                        if m2:
                            # 组合行同样过内容校验(短标题/量词行不组合成幽灵条文);
                            # rest 为空(独立数字行排版 "5."+下一行 "63")是真条文号,保留
                            rest2 = lines[j].translate(OCR_NORM_TABLE)[m2.end():].lstrip(' \t　\xa0')
                            if rest2 and (rest2[0].isdigit() or rest2[0] in '米倍条次天年％%'):
                                break
                            if rest2 and not re.search(r'[，。；：、]', rest2) \
                                    and len(re.sub(r'[\s　\xa0]+', '', rest2)) <= 6:
                                break
                            no = f"{m_solo.group(1)}.{m2.group(1)}"
                            printed = p - offset if offset else ''
                            page_rows.append((no, str(p), str(printed),
                                              '1' if (expl and p >= expl) else '0',
                                              '1' if meta["type"] == "ocr" else '0',
                                              'low', lines[j].strip()[:40].replace('\t', ' ')))
                            break
                continue
            no = '.'.join(g for g in m.groups() if g)
            printed = p - offset if offset else ''
            # 注意: 独立数字行排版("4.8.2" 单独一行)时,同页下一行可能是
            # 双栏流式提取出的其他条文残留,不可抓取以免张冠李戴
            text = s[:40].replace('\t', ' ')
            page_rows.append((no, str(p), str(printed),
                              '1' if (expl and p >= expl) else '0',
                              '1' if meta["type"] == "ocr" else '0',
                              '', text))
        if len(page_rows) > 40:
            # 整页疑似表格/术语表: 该页条目标记 low_confidence
            page_rows = [r[:5] + ('low',) + (r[6],) for r in page_rows]
        rows.extend(page_rows)
    return rows


def write_clauses(book_dir, rows):
    lines = ["条文号\tPDF页\t正文页\texpl\tocr\t低置信\t行首原文"]
    lines.extend("\t".join(str(x) for x in r) for r in rows)
    atomic_write_text(Path(book_dir) / "clauses.idx", "\n".join(lines) + "\n")


# ---------- 子命令 ----------

def _cmd_index(args, cfg):
    data_dir = cfg["data_dir"]
    Path(data_dir).mkdir(parents=True, exist_ok=True)
    shelf = load_shelf(data_dir)
    chars = load_chars(cfg["common_chars"])
    only = {e.lower().lstrip(".") for e in (args.only or [])}
    if args.all:
        files_map, _hints = scan_library(Path(cfg["library_dir"]))
        if not files_map:
            die(f"库目录没有支持的文件: {cfg['library_dir']}")
        sources = [Path(v) for v in files_map.values()]
        if only:
            sources = [p for p in sources if p.suffix.lower().lstrip(".") in only]
    else:
        sources = [Path(p) for p in args.pdfs]
        for p in sources:
            if not p.exists():
                die(f"文件不存在: {p}")
        if only:
            sources = [p for p in sources if p.suffix.lower().lstrip(".") in only]
    if not sources:
        die("请指定源文件路径,或用 --all 处理库目录全部支持的文件")
    # 必须在并发前冻结分配；_index_one 只读此映射。
    cfg["_book_ids"] = plan_book_ids(sources, cfg, shelf)
    print(f"待处理 {len(sources)} 个文件,并行度 {args.jobs}")
    global _COM_CTX
    _COM_CTX = ComContext()

    def _needs_com(p):
        """魔数分流: xlsx/docx 非 zip(伪格式/加密)→ COM;doc/xls 是 zip(改名)→ 纯 Python。"""
        fmt = p.suffix.lower().lstrip(".")
        if fmt in ("docx", "xlsx") and not _is_zip(p):
            return True
        if fmt in ("doc", "xls") and _is_zip(p):
            return False
        return fmt in _COM_FMTS

    results = []
    pure = [p for p in sources if not _needs_com(p)]
    com = [p for p in sources if _needs_com(p)]
    # 纯 Python 格式进线程池;COM 格式串行(apartment 模型,严禁进线程池)
    if args.jobs > 1 and len(pure) > 1:
        with ThreadPoolExecutor(max_workers=args.jobs) as ex:
            results = list(ex.map(lambda p: _index_one(p, cfg, chars), pure))
    else:
        results = [_index_one(p, cfg, chars) for p in pure]
    for i, p in enumerate(com, 1):
        print(f"[COM {i}/{len(com)} 串行] {p.name}", flush=True)
        results.append(_index_one(p, cfg, chars))
    _COM_CTX.quit_all()
    _COM_CTX = None
    # 书架合并必须在单线程做(_index_one 并行时不写 bookshelf.json)
    shelf = load_shelf(data_dir)
    ok = skip = fail = 0
    for path, status, msg, entry in results:
        if status == "ok":
            ok += 1
            old = next((b for b in shelf["books"] if b["id"] == entry["id"]), None)
            if old:
                for k in ("category", "alias", "replaces", "replaced_by", "note"):
                    if old.get(k):
                        entry[k] = old[k]
                shelf["books"] = [b for b in shelf["books"] if b["id"] != entry["id"]]
            shelf["books"].append(entry)
        elif status == "skip":
            skip += 1
        else:
            fail += 1
    shelf["books"].sort(key=lambda b: b["id"])
    save_shelf(data_dir, shelf)
    for path, status, msg, _ in results:
        mark = {"ok": "✔", "skip": "·", "fail": "✘"}[status]
        print(f"  {mark} {Path(path).name}: {msg}")
    print(f"完成: 新增/更新 {ok} 本,跳过 {skip} 本,失败 {fail} 本")
    if fail:
        sys.exit(1)


def cmd_index(args, cfg):
    with library_lock(cfg["data_dir"]):
        _cmd_index(args, cfg)


def _index_one(src_path, cfg, chars):
    src_path = Path(src_path)
    try:
        data_dir = cfg["data_dir"]
        shelf = load_shelf(data_dir)
        title, std_no, version, seq = parse_filename(src_path.name)
        sig = source_signature(src_path)
        mtime = sig["display_mtime"]
        # file 记录相对 library_dir 的路径(guifansrc 可多层文件夹),book_id = 源文件名去扩展名
        rel = _rel_source(src_path, cfg)
        fmt = src_path.suffix.lower().lstrip(".") or "unknown"
        if fmt not in {e.lstrip(".") for e in INDEXABLE_EXTS}:
            return src_path, "fail", f"不支持的扩展名: .{fmt}(支持: {', '.join(sorted(INDEXABLE_EXTS))})", None
        base_id = Path(rel).stem.rstrip(' .')  # 去尾空格/点(dup 碰撞检测用全名,截断见下)
        old = next((b for b in shelf["books"] if b["file"] == rel), None)
        if old:
            std_no = old.get("std_no") or std_no
            title = old.get("title") or title
            version = old.get("version") or version
        # book_id: 源文件名;与旧书架 id 兼容(迁移后 id=文件名);同 stem 不同路径时加父目录前缀
        book_id = cfg.get("_book_ids", {}).get(rel) or (old["id"] if old else _short_id(base_id))
        source_changed = old is not None and not entry_matches_source(old, sig)
        if old and not cfg.get("force") and not cfg.get("rebuild") and old.get("status") == "indexed" and not source_changed:
            # 迁移 gap 防御: 书架已登记但新布局路径无索引数据(旧布局书永久不可达)。
            # 旧布局 data_dir/<book_id>/ 数据(含 OCR 缓存)存在 → 移到新路径复用,
            # 避免整本重 OCR;否则才重索引
            bdir_key = {"file": rel, "id": book_id}
            new_bd = Path(book_data_dir(cfg, bdir_key))
            if (new_bd / "meta.json").exists():
                return src_path, "skip", f"已是最新索引(book_id={book_id})", None
            old_bd = Path(data_dir) / book_id
            if old_bd.exists():
                import shutil
                shutil.move(str(old_bd), str(new_bd))
                print(f"[{book_id}] 迁移旧布局数据 {old_bd} → {new_bd}", flush=True)
            print(f"[{book_id}] 书架已登记但新路径无索引数据,重索引", flush=True)
        # 索引数据目录与 guifansrc 结构一致: data_dir / 源文件相对父目录 / book_id
        book_dir = Path(book_data_dir(cfg, {"file": rel, "id": book_id}))
        book_dir.mkdir(parents=True, exist_ok=True)
        # 按格式分派提取: pdf=现有全流程(probe→text/ocr);图片=整图 OCR;
        # ofd=物理页文本层直取;其余=文本行→虚拟分页(COM 格式先另存 docx/xlsx)
        pages = 0
        probe = None
        texts = {}
        pages_dir = None
        meta_ocr = {"failed_pages": [], "pages_done": 0}
        note_low = ""
        if fmt == "pdf":
            probe = probe_pdf(src_path, chars)
            doc = fitz.open(src_path)
            pages = doc.page_count
            if probe["type"] == "ocr":
                print(f"[{book_id}] 扫描/乱码书,整本 OCR(共 {pages} 页)...", flush=True)
                failed, texts = _ocr_book(src_path, book_dir, cfg, pages, book_id, "pdf",
                                           force=bool(cfg.get("force") or source_changed))
                meta_ocr = {"failed_pages": failed, "pages_done": pages - len(failed)}
                pages_dir = "ocr"
            else:
                texts = extract_text_pages(doc, book_dir, pages)  # 文字书内存提取,不落盘
                # 全文质量二次检测: probe 抽样可能漏判伪文字版(部分页 ToUnicode CMap 损坏,
                # 提取为控制字符或 CJK 扩展区乱码,如 139/140;三项任一显著 → 转整本 OCR)
                full = ''.join(texts.values())
                nonws = [c for c in full if not c.isspace()]
                cjk = [c for c in nonws if is_cjk(c)]
                full_cov = len([c for c in cjk if c in chars]) / len(cjk) if cjk else 0
                ext_ratio = len([c for c in cjk if not (0x4E00 <= ord(c) <= 0x9FFF)]) / len(cjk) if cjk else 0
                ctrl_ratio = len([c for c in nonws if ord(c) < 32 or ord(c) == 0xFFFD]) / len(nonws) if nonws else 0
                # 无 CJK 的书(管理文件/英文手册): 覆盖率恒 0,不得触发 OCR 判定;
                # ctrl_ratio 独立于 cjk,控制字符泛滥仍转 OCR
                if (cjk and (full_cov < 0.92 or ext_ratio > 0.15)) or (nonws and ctrl_ratio > 0.05):
                    print(f"[{book_id}] 全文质量二次检测不过(覆盖率{full_cov:.3f}/扩展区{ext_ratio:.3f}/控制字符{ctrl_ratio:.3f}),转整本 OCR...", flush=True)
                    probe = dict(probe, type="ocr", rule=probe.get("rule", []) + ["rescued-by-fulltext-check"])
                    failed, texts = _ocr_book(src_path, book_dir, cfg, pages, book_id, "pdf",
                                               force=bool(cfg.get("force") or source_changed))
                    meta_ocr = {"failed_pages": failed, "pages_done": pages - len(failed)}
                    pages_dir = "ocr"
            doc.close()
        elif fmt in _IMAGE_FMTS:
            # 图片书: 每图/每帧 = 1 页,tesseract 整图 OCR(与 PDF OCR 书同构)
            pages = image_page_count(src_path)
            probe = {"type": "ocr", "pages": pages, "rule": ["image-source"],
                     "cov_median": None, "punct_median": None,
                     "empty_ratio": 1, "bad_ratio": 1, "clause_hits": 0}
            print(f"[{book_id}] 图片书,整本 OCR(共 {pages} 页)...", flush=True)
            failed, texts = _ocr_book(src_path, book_dir, cfg, pages, book_id, fmt,
                                       force=bool(cfg.get("force") or source_changed))
            meta_ocr = {"failed_pages": failed, "pages_done": pages - len(failed)}
            pages_dir = "ocr"
        elif fmt == "ofd":
            # OFD 有物理页: Content.xml 文本层直取(纯 Python),不虚拟分页
            page_texts = extract_ofd(src_path)
            pages = len(page_texts)
            texts = {i + 1: t for i, t in enumerate(page_texts)}
            probe = probe_text(texts, chars)
            pages_dir = "extracted"
            _write_extracted(book_dir, texts)
        else:
            # 虚拟分页文本书(含 COM 转换): 行列表 → 虚拟页 → 落盘 extracted/
            if fmt == "docx":
                rows = extract_docx(src_path) if _is_zip(src_path) \
                    else extract_word_com(src_path, cfg, _COM_CTX)
            elif fmt == "xlsx":
                rows = extract_xlsx(src_path) if _is_zip(src_path) \
                    else extract_excel_com(src_path, cfg, _COM_CTX)
            elif fmt in ("txt", "md"):
                rows = extract_txt(src_path)
            elif fmt in _COM_FMTS:
                if fmt == "doc" and _is_zip(src_path):
                    rows = extract_docx(src_path)   # 扩展名 .doc 实为 docx(改名文件)
                elif fmt == "xls" and _is_zip(src_path):
                    rows = extract_xlsx(src_path)   # 扩展名 .xls 实为 xlsx
                else:
                    rows = (extract_word_com(src_path, cfg, _COM_CTX)
                            if fmt in ("doc", "wps") else extract_excel_com(src_path, cfg, _COM_CTX))
            else:
                return src_path, "fail", f"不支持的扩展名: .{fmt}", None
            pages_text = virtual_paginate(rows)
            pages = len(pages_text)
            texts = {i + 1: t for i, t in enumerate(pages_text)}
            probe = probe_text(texts, chars)
            pages_dir = "extracted"
            _write_extracted(book_dir, texts)
        meta = {
            "id": book_id, "title": title, "std_no": std_no, "version": version,
            # file 用相对 library_dir 路径,与书架条目一致(book_data_dir 的入参形态,
            # 防止未来 book_data_dir(cfg, meta) 解析到错误目录)
            "file": rel, "source_abs": str(src_path), "pdf_abs": str(src_path),
            "pages": pages, "type": probe["type"], "probe": probe,
            "source_mtime": mtime, "pdf_mtime": mtime,  # 双写兼容期
            "source_mtime_ns": sig["mtime_ns"], "source_size": sig["size"],
            "fmt": fmt, "virtual": fmt in _VIRTUAL_FMTS,
            "pages_dir": pages_dir, "ocr": meta_ocr,
            "indexed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        if not any(t.strip() for t in texts.values()):
            note_low = "文档无可提取文本(空/纯图)(low_confidence)"
            meta["note"] = note_low
        missing = [p for p in range(1, pages + 1) if p not in texts]
        if missing:
            meta["status"] = "indexing"
            atomic_write_text(book_dir / "meta.json", json.dumps(meta, ensure_ascii=False, indent=2) + "\n")
            return src_path, "fail", f"缺 {len(missing)} 页文本,OCR/提取未完成,status=indexing,重跑 index 续跑", None
        # TOC 三路径
        toc_zone = find_toc_zone(texts, pages)
        zone_end = toc_zone[1]
        entries, offset, toc_source = None, None, None
        if probe["type"] == "text" and fmt == "pdf":
            d = fitz.open(src_path)
            bm = toc_from_bookmarks(d)
            d.close()
            if bm:
                entries, toc_source = _bm_entries(bm), "bookmark"
        if entries is None:
            r = toc_from_tocpages(texts, pages)
            if r:
                entries, zone_end, offset = r
                toc_source = "parsed"
        expl_start = detect_expl_start(texts, pages, zone_end)
        if entries is None:
            scanned, expl2 = toc_from_bodyscan(texts, pages, zone_end + 1)
            if scanned:
                entries, toc_source = scanned, "body_scan"
                if expl_start is None:
                    expl_start = expl2
        # OCR 书: 标题常被误读但数字识别可靠 → 条文号聚类定页,
        # 无条文的章(范围/引用文件等)按章号顺序在相邻区间找"章号+短标题"行补
        if probe["type"] == "ocr":
            cluster = _clause_prefix_first_pages(texts, zone_end, pages)
            merged_ch = {}
            for n, p in sorted(cluster.items()):
                if 1 <= n <= 40:  # 过滤 OCR 误读产生的超大章号
                    merged_ch[str(n)] = (_heading_on_page(texts, p, str(n)) or f'第{n}章', p)
            prev_page = zone_end
            sane = [n for n in cluster if 1 <= n <= 40]
            max_n = max(sane) if sane else 0
            for n in range(1, min(max_n, 40) + 1):
                if str(n) in merged_ch:
                    prev_page = max(prev_page, merged_ch[str(n)][1])
                    continue
                for p in range(prev_page, min(prev_page + 16, pages) + 1):
                    ch_t = _heading_on_page(texts, p, str(n))
                    if ch_t and len(re.sub(r'[\s　\xa0]+', '', ch_t)) <= 15 \
                            and '。' not in ch_t:
                        merged_ch[str(n)] = (ch_t, p)
                        prev_page = p
                        break
            if len(merged_ch) >= 2:
                entries = [(l, t, p) for l, (t, p) in merged_ch.items()]
                entries.sort(key=lambda e: e[2])
                toc_source = "clause_cluster"
        # 条文号聚类补章: 文字书解析/扫描漏掉的章,用"第 N 章条文首次出现的页"补齐
        if entries and probe["type"] == "text":
            first_pages = _clause_prefix_first_pages(texts, zone_end, pages)
            have = {int(e[0]) for e in entries if e[0].isdigit()}
            for n in sorted(first_pages):
                # 章号上界与 OCR 路径一致(1-40),防 "58 . 2 …" 幽灵章号补出伪章
                if n < 1 or n > 40 or n in have or first_pages[n] <= zone_end:
                    continue
                # 条文说明区(expl 后)的条文号不补章;占位标题保证章边界(标题行
                # 缺失的真章不丢,如 015/065);同页 ≥2 个占位章再丢弃(171 伪章)
                if expl_start and first_pages[n] >= expl_start:
                    continue
                ch_title = _heading_on_page(texts, first_pages[n], str(n)) or f'第{n}章'
                entries.append((str(n), ch_title, first_pages[n]))
        # 同页 ≥2 个占位章(找不到标题) → 编号列表噪声(条文说明区引用行),丢弃该页占位章
        if entries:
            ph_pages = Counter(e[2] for e in entries if re.match(r'^第\d{1,2}章$', e[1]))
            entries = [e for e in entries
                       if not (re.match(r'^第\d{1,2}章$', e[1]) and ph_pages.get(e[2], 0) >= 2)]
        # 条文说明 fallback(仅文字书): 目录条目须晚于所有正文章,防目录页噪声劫持
        if probe["type"] == "text" and expl_start is None and entries:
            ep = next((pg for _, t, pg in entries if '条文说明' in t), None)
            if ep and ep >= max((pg for _, _, pg in entries), default=0):
                expl_start = ep
        meta.update({"toc_zone_end": zone_end, "offset": offset,
                     "toc_source": toc_source or "none", "expl_start": expl_start})
        # 切章与条文
        chapters = None
        if entries:
            chapters = build_chapters(entries, expl_start, pages)
        else:
            # 降级: 按 ~30 页均分,保证可导航
            chunk = 30
            chapters = [{"label": f"第{i}部分", "title": f"第{i}部分(无目录,按页均分)",
                         "start": s, "is_expl": False}
                        for i, s in enumerate(range(zone_end + 1, pages + 1, chunk), 1)]
            meta["note"] = "未解析到目录,章节按 30 页均分(low_confidence)"
        idx_rows = build_clauses(texts, meta, offset)
        write_clauses(book_dir, idx_rows)
        write_toc_md(book_dir, meta, chapters)
        written = write_chapters(book_dir, texts, meta, chapters, idx_rows)
        meta.update({"chapter_list": [{"no": i, "label": c["label"], "title": c["title"],
                                       "start": c["start"], "is_expl": c["is_expl"]}
                                      for i, c in enumerate(chapters, 1)],
                     "clause_count": len(idx_rows),
                     "status": "indexed", "note": meta.get("note", "")})
        atomic_write_text(book_dir / "meta.json", json.dumps(meta, ensure_ascii=False, indent=2) + "\n")
        # 更新书架(保留人工字段)
        new_entry = {
            "id": book_id, "title": title, "std_no": std_no, "file": rel,
            "category": "", "pages": pages, "type": probe["type"],
            "status": "indexed", "toc_source": meta["toc_source"], "version": version,
            "alias": [], "replaces": [], "replaced_by": None,
            "ocr": meta["ocr"], "indexed_at": meta["indexed_at"],
            "source_mtime": mtime, "pdf_mtime": mtime,  # 双写兼容期
            "source_mtime_ns": sig["mtime_ns"], "source_size": sig["size"],
            "fmt": fmt, "probe": probe, "note": meta.get("note", ""),
        }
        return src_path, "ok", (f"book_id={book_id} type={probe['type']} fmt={fmt} "
                                f"toc={meta['toc_source']} 章数={len(chapters)} "
                                f"条文数={len(idx_rows)}"), new_entry
    except Exception as e:
        return src_path, "fail", str(e), None


def cmd_list(args, cfg):
    shelf = load_shelf(cfg["data_dir"])
    books = shelf.get("books", [])
    if args.type:
        books = [b for b in books if b["type"] == args.type]
    if args.category:
        books = [b for b in books if args.category in (b.get("category") or "")]
    if args.query:
        q = args.query.lower()
        books = [b for b in books if q in b["id"].lower() or q in b["title"].lower()
                 or q in norm_no(b.get("std_no", "")) or any(q in a.lower() for a in b.get("alias", []))]
    if not books:
        print("(书架为空)")
        return
    for b in books:
        cc = "?"
        meta_f = book_data_dir(cfg, b) / "meta.json"
        if meta_f.exists():
            cc = json.loads(meta_f.read_text(encoding="utf-8")).get("clause_count", "?")
        repl = f" →已由 {b['replaced_by']} 替代" if b.get("replaced_by") else ""
        print(f"{b['id']}\t{b.get('std_no','')}\t{b['title']}\t{b['type']}\t{b.get('fmt','pdf')}"
              f"\t{b['pages']}页\t{b['status']}\t{b.get('category') or '-'}\t条文{cc}{repl}")


def cmd_toc(args, cfg):
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    f = book_data_dir(cfg, b) / "toc.md"
    if not f.exists():
        die(f"{b['id']} 没有 toc.md,先 `spec.py index`")
    lines = f.read_text(encoding="utf-8").splitlines()
    if args.chapter:
        for ln in lines:
            if ln.startswith(f"| {args.chapter} "):
                print(ln)
                return
        die(f"第 {args.chapter} 章不在目录中")
    print("\n".join(lines))


def cmd_clause(args, cfg):
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    f = book_data_dir(cfg, b) / "clauses.idx"
    if not f.exists():
        die(f"{b['id']} 没有 clauses.idx,先 `spec.py index`")
    rows = [ln.split("\t") for ln in f.read_text(encoding="utf-8").splitlines()[1:] if ln.strip()]
    # 查询词去空白归一化(从章文件复制 "4. 2. 12" 直查也能命中归一化条文号索引)
    q = norm_no(args.clause).translate(OCR_NORM_TABLE)
    hits = [r for r in rows if r[0].lower() == q]
    if hits:
        for r in hits:
            expl = "(条文说明)" if r[3] == '1' else ""
            low = "(低置信)" if len(r) > 5 and r[5] else ""
            printed = f"|正文页{r[2]}" if r[2] else ""
            print(f"{r[0]}  PDF页{r[1]}{printed} {expl}{low}")
            if r[-1] == r[0]:
                print(f"   (独立条文号行,内容见下页 — 用 read {b['id']} {r[1]} 或 grep 关键词确认)")
            else:
                print(f"   {r[-1]}")
        return
    # 未命中: 相邻条文提示
    def key(no):
        try:
            return tuple(int(x) for x in no.split('.'))
        except ValueError:
            return (10 ** 9,)

    try:
        qk = tuple(int(x) for x in q.split('.'))
    except ValueError:
        qk = (10 ** 9,)

    def dist(n):
        a, b = key(n) + (0, 0, 0), qk + (0, 0, 0)
        return sum(abs(x - y) * w for x, y, w in zip(a[:3], b[:3], (10000, 100, 1)))

    all_nos = sorted({r[0] for r in rows}, key=key)
    near = sorted(all_nos, key=lambda n: (dist(n), key(n)))[:args.near]
    print(f"未命中条文 {args.clause}。相邻条文号:")
    for n in near:
        r = next(r for r in rows if r[0] == n)
        print(f"  {n}(PDF页{r[1]}) {r[-1]}")
    if b["type"] == "ocr":
        print("提示: 本书为 OCR 书,条文号可能被误读,建议 `spec.py grep <book> '<关键词>'` 在页内确认")


def cmd_read(args, cfg):
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    end = args.end or args.page
    if end - args.page + 1 > 20:
        die("单次最多 20 页,请分批")
    if args.page < 1 or end > b["pages"]:
        die(f"页码超界: 本书共 {b['pages']} 页")
    book_dir = book_data_dir(cfg, b)
    meta_local = {}
    meta_f = book_dir / "meta.json"
    if meta_f.exists():
        meta_local = json.loads(meta_f.read_text(encoding="utf-8"))
    offset = meta_local.get("offset")
    is_ocr = b["type"] == "ocr"
    virtual = meta_local.get("virtual")
    # 文本来源分派: pages_dir=ocr(OCR 书/图片书) / extracted(非 PDF 文本落盘) /
    # 缺省(PDF 文字书现场提取)
    pages_dir = meta_local.get("pages_dir")
    src_abs = None
    m = _src_abs(meta_local)
    if m and Path(m).exists():
        src_abs = m
    elif not pages_dir and not is_ocr:
        src_abs = _find_source(cfg, b)   # 找不到时内部 die 报错
    chars = load_chars(cfg["common_chars"])
    doc = None if (pages_dir or is_ocr) else fitz.open(src_abs)
    warn = False
    try:
        for p in range(args.page, end + 1):
            if pages_dir == "ocr":
                f = book_dir / "ocr" / f"{p:03d}.txt"
                if not f.exists():
                    die(f"ocr/{p:03d}.txt 缺失 — OCR 未完成,先 `spec.py ocr {b['id']}` 或重跑 index")
                t = f.read_text(encoding="utf-8")
            elif pages_dir == "extracted":
                f = book_dir / "extracted" / f"{p:03d}.txt"
                if not f.exists():
                    die(f"extracted/{p:03d}.txt 缺失 — 提取未完成,重跑 index")
                t = f.read_text(encoding="utf-8")
            else:
                t = doc[p - 1].get_text()  # PDF 文字书现场提取,不落盘
            printed = f"|正文页{p - offset}" if offset else ""
            marker = f"【第 {p} 页{printed}】" + ("【OCR】" if is_ocr else "") \
                + ("【虚拟页】" if virtual else "")
            print(f"===== {marker} =====")
            print(t.strip())
            # 质量自检: 疑似乱码页提示勿引用
            s = page_stats(t, chars)
            if s["n"] >= 20 and ((s["cov"] or 0) < 0.95 or (s["punct"] or 0) > 0.08):
                warn = True
    finally:
        if doc:
            doc.close()
    if warn:
        print("[警告] 部分页疑似乱码,勿作为原文引用,建议 `spec.py ocr` 该页")


def _find_source(cfg, b):
    p = Path(cfg["library_dir"]) / b["file"]
    if p.exists():
        return str(p)
    # 兼容旧记录(纯文件名,无子目录)
    for q in Path(cfg["library_dir"]).rglob(b["file"]):
        return str(q)
    die(f"库目录中找不到源文件: {b['file']}")


def cmd_grep(args, cfg):
    if not args.all_books and not args.book:
        die("需要 <book> 或用 --all 跨全部已索引书")
    shelf = load_shelf(cfg["data_dir"])
    pattern = re.compile(args.pattern, re.I if args.ignore_case else 0)
    targets = []
    if args.all_books:
        for b in shelf.get("books", []):
            if b.get("status") == "indexed":
                targets.append((b["id"], book_data_dir(cfg, b) / "chapters"))
    else:
        b = find_book(shelf, args.book)
        ch = book_data_dir(cfg, b) / "chapters"
        if not ch.exists():
            die(f"{b['id']} 无章节文件,先 `spec.py index`")
        targets.append((b["id"], ch))
    count = 0
    # OCR 文本字间常带空格("隧 道"),压缩版正则做二次匹配
    pat_c = re.compile(re.sub(r'[\s　\xa0]+', '', args.pattern),
                       re.I if args.ignore_case else 0)

    def _c(s):
        return re.sub(r'[\s　\xa0]+', '', s)

    for bid, chdir in targets:
        for f in sorted(list(chdir.glob("*.md")) + list(chdir.glob("*.txt"))):
            lines = f.read_text(encoding="utf-8").splitlines()
            page = ""
            for i, ln in enumerate(lines):
                m = re.match(r'^【第 (\d+) 页', ln)
                if m:
                    page = m.group(1)
                if pattern.search(ln) or pat_c.search(_c(ln)):
                    lo, hi = max(0, i - args.ctx), min(len(lines), i + args.ctx + 1)
                    for j in range(lo, hi):
                        print(f"{bid}/{f.name}:{j + 1} [第{page}页] {lines[j]}")
                    print("---")
                    count += 1
                    if args.max and count >= args.max:
                        print(f"(达到上限 {args.max} 条)")
                        return
    if count == 0:
        print(f"无匹配: {args.pattern}")


def cmd_ocr(args, cfg):
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    fmt = b.get("fmt", "pdf")
    if fmt not in ("pdf",) and fmt not in _IMAGE_FMTS:
        die(f"本书为文字格式(fmt={fmt}),无 OCR 需求")
    src = _find_source(cfg, b)
    # 参数 clamp 到 [1, pages](防 fitz 越界报错);start>end 报错(防静默空跑)。
    # end 不做 max(1): 负端(--end -3)留给 start>end 守卫报错,不静默成第 1 页;
    # 显式 0 是 typo(--end 0 会 or 成全书)→ 报错而非静默整本 OCR
    if args.start == 0 or args.end == 0:
        die("页码从 1 开始,--start/--end 不能为 0")
    start = max(1, min(args.start or 1, b["pages"]))
    end = min(args.end or b["pages"], b["pages"])
    if start > end:
        die(f"--start({start}) 大于 --end({end}),OCR 范围无效")
    print(f"OCR {b['id']} 第 {start}-{end} 页(共 {b['pages']} 页)...", flush=True)
    failed = run_ocr_pages(src, book_data_dir(cfg, b), cfg, start, end, args.force, fmt)
    if failed:
        print(f"失败 {len(failed)} 页: {failed}")
        sys.exit(1)
    print("OCR 完成")


def cmd_img(args, cfg):
    """渲染书源 PDF 的一页(或范围)为 PNG,供 agent 视觉复核/看图。
    输出写入 <book_dir>/recheck/pNNN.png,并把路径打印到 stdout。读只操作,不改索引。"""
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    if b.get("fmt", "pdf") != "pdf":
        die(f"本书 fmt={b.get('fmt')},非 PDF,无法渲染页面")
    src = _find_source(cfg, b)
    end = args.end or args.page
    if args.page < 1 or end > b["pages"] or args.page > end:
        die(f"页码超界: 本书共 {b['pages']} 页;要求 1<=start<=end<=pages")
    book_dir = book_data_dir(cfg, b)
    recheck_dir = book_dir / "recheck"
    recheck_dir.mkdir(exist_ok=True)
    doc = fitz.open(src)
    out = []
    try:
        for p in range(args.page, end + 1):
            pix = doc.load_page(p - 1).get_pixmap(dpi=args.dpi)
            png = recheck_dir / f"p{p:03d}.png"
            pix.save(str(png))
            out.append(str(png))
    finally:
        doc.close()
    print("\n".join(out))


def cmd_recheck(args, cfg):
    """报告一本书的低置信/失败页(供 agent 决定视觉复核哪些页)。只读。"""
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    book_dir = book_data_dir(cfg, b)
    meta = {}
    mf = book_dir / "meta.json"
    if mf.exists():
        meta = json.loads(mf.read_text(encoding="utf-8"))
    print(f"book_id={b['id']} pages={b.get('pages')} type={b.get('type')} pages_dir={meta.get('pages_dir')} toc={b.get('toc_source')}")
    failed = (meta.get("ocr") or {}).get("failed_pages") or []
    print(f"failed_pages: {failed if failed else '(无)'}")
    if meta.get("note"):
        print(f"note: {meta['note']}")
    idx = book_dir / "clauses.idx"
    low = []
    if idx.exists():
        lines = idx.read_text(encoding="utf-8").splitlines()
        for line in lines[1:]:
            cols = line.split("\t")
            if len(cols) > 5 and cols[5].strip():
                low.append((cols[0], cols[1], cols[5]))
    print(f"low_confidence_clauses: {len(low)}")
    for clause, pdfp, marker in low[:args.max]:
        print(f"  {clause}\tp{pdfp}\t{marker}")
    if len(low) > args.max:
        print(f"  ...(余 {len(low) - args.max} 条)")
    if not failed and not low and not meta.get("note"):
        print("未发现问题页(无需视觉复核)")


def cmd_set_page(args, cfg):
    """回写某页被修正的文本(ocr/NNN.txt 或 extracted/NNN.txt),供 agent 视觉修正后落盘。
    配合 `index <源文件> --rebuild` 重新派生章/条文。仅限有落盘页文本的书(OCR 书/非 PDF 文本书)。"""
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    book_dir = book_data_dir(cfg, b)
    meta = {}
    mf = book_dir / "meta.json"
    if mf.exists():
        meta = json.loads(mf.read_text(encoding="utf-8"))
    pages_dir = meta.get("pages_dir")
    if pages_dir not in ("ocr", "extracted"):
        die("本书无落盘页文本(纯 PDF 文字书按需现场提取),不支持 set-page;仅 OCR 书/非 PDF 文本书可回写")
    page = args.page
    if page < 1 or page > b["pages"]:
        die(f"页码超界: 本书共 {b['pages']} 页")
    text = args.text if args.text is not None else sys.stdin.read()
    target = book_dir / pages_dir / f"{page:03d}.txt"
    target.parent.mkdir(parents=True, exist_ok=True)
    atomic_write_text(target, text)
    src_rel = b["file"]
    full = Path(cfg["library_dir"]) / src_rel
    print(f"wrote {target} ({len(text)} chars)")
    print(f"NEXT: `spec.py index \"{full}\" --rebuild` 重新派生章/条文/目录")


def cmd_status(args, cfg):
    shelf = load_shelf(cfg["data_dir"])
    # 键用相对 library_dir 路径(guifansrc 可多层文件夹,同名文件不混淆)
    lib, hints = scan_library(Path(cfg["library_dir"]))
    books = shelf.get("books", [])
    indexed_names = {b["file"] for b in books}
    new_sources = sorted(set(lib) - indexed_names)
    missing = [b for b in books if b["file"] not in lib]
    changed = []
    for b in books:
        src = lib.get(b.get("file"))
        if src:
            try:
                if not entry_matches_source(b, source_signature(src)):
                    changed.append((b, src))
            except OSError as e:
                print(f"[警告] 无法读取 {b['file']} 的文件属性: {e}")
    # 疑似换版: 新文件归一化名与已索引书名相同
    def norm_name(x):
        x = re.sub(r'[（(][^（）()]*[）)]', '', x)
        x = re.sub(r'[\s.+—\-_]', '', x.lower())
        return re.sub(r'\d{4}', '', x)
    title_map = {norm_name(b["title"] + b.get("std_no", "")): b["id"] for b in books}
    print("== 库一致性检查 ==")
    if not new_sources and not missing and not changed:
        print("一致: 库目录与书架索引同步")
    for n in new_sources:
        hint = ""
        key = norm_name(Path(n).stem)
        for tk, bid in title_map.items():
            if key and (key == tk or key in tk or tk in key):
                hint = f"  疑似换版: 与已索引 {bid} 名称高度相似,确认后 `spec.py index` 并标记替代关系"
                break
        print(f"[新增] {n} 未建索引 → `spec.py index \"{lib[n]}\"{hint}")
    for b in missing:
        print(f"[缺失] {b['id']}({b['file']}) 源文件已不在库目录 → `spec.py remove {b['id']}` 清理")
    for b, src in changed:
        print(f"[更新] {b['file']} 源文件已变更 → `spec.py index \"{src}\"`")
    for name, why in hints:
        print(f"[未索引] {name}({why})")
    print("")
    print("== 书架健康 ==")
    for b in books:
        meta_f = book_data_dir(cfg, b) / "meta.json"
        if meta_f.exists():
            m = json.loads(meta_f.read_text(encoding="utf-8"))
            n_ch = len(list((book_data_dir(cfg, b) / "chapters").glob("*.md"))
                      + list((book_data_dir(cfg, b) / "chapters").glob("*.txt")))
            ocr_failed = m.get("ocr", {}).get("failed_pages", [])
            repl = f" 已被 {b['replaced_by']} 替代" if b.get("replaced_by") else ""
            print(f"{b['id']}\t{b['type']}\t{b['status']}\t{b.get('fmt', 'pdf')}\t"
                  f"toc={m.get('toc_source')}\t章文件={n_ch}\t条文={m.get('clause_count')}\t"
                  f"OCR失败页={len(ocr_failed)}{repl}")
        else:
            print(f"{b['id']}\t{b.get('type','?')}\t{b.get('status','?')}\t(无 meta.json)")


def _cmd_remove(args, cfg):
    shelf = load_shelf(cfg["data_dir"])
    b = find_book(shelf, args.book)
    if args.mark_superseded:
        if args.mark_superseded == b["id"]:
            die("替代目标不能是该书自身")
        tgt = next((x for x in shelf["books"] if x["id"] == args.mark_superseded), None)
        if not tgt:
            die(f"替代目标不存在: {args.mark_superseded}。请先 index 新版本，再标记替代关系")
        b["status"] = "superseded"
        b["replaced_by"] = args.mark_superseded
        if b["id"] not in tgt.get("replaces", []):
            tgt.setdefault("replaces", []).append(b["id"])
        save_shelf(cfg["data_dir"], shelf)
        print(f"{b['id']} 已标记为 superseded,由 {args.mark_superseded} 替代(索引保留,历史版本仍可查)")
        return
    import shutil
    d = book_data_dir(cfg, b)
    if d.exists():
        shutil.rmtree(d)
    shelf["books"] = [x for x in shelf["books"] if x["id"] != b["id"]]
    save_shelf(cfg["data_dir"], shelf)
    print(f"已删除 {b['id']} 的索引与书架登记(源 PDF 未动)")


def cmd_remove(args, cfg):
    with library_lock(cfg["data_dir"]):
        _cmd_remove(args, cfg)


def cmd_update_chars(args, cfg):
    """生成/刷新常用字表: 从干净文字版 PDF 统计字频,取前 3500 字(仅 PDF 参与)。"""
    pdfs = [Path(p) for p in (args.from_pdfs or [])]
    if not pdfs:
        shelf = load_shelf(cfg["data_dir"])
        for b in shelf.get("books", []):
            if b.get("type") == "text" and b.get("fmt", "pdf") == "pdf":
                p = _src_abs(b) or _find_source(cfg, b)
                if Path(p).exists():
                    pdfs.append(Path(p))
    if not pdfs:
        die("没有来源 PDF: 用 --from-pdfs <干净文字版PDF...> 指定")
    counts = Counter()
    for p in pdfs:
        doc = fitz.open(p)
        for i in range(doc.page_count):
            for ch in doc[i].get_text():
                if is_cjk(ch):
                    counts[ch] += 1
        doc.close()
        print(f"  统计 {p.name}: 累计 {len(counts)} 字")
    top = [ch for ch, _ in counts.most_common(3500)]
    out = SKILL_DIR / cfg["common_chars"]
    out.parent.mkdir(parents=True, exist_ok=True)
    atomic_write_text(out, "\n".join("".join(top[i:i + 50]) for i in range(0, len(top), 50)) + "\n")
    print(f"常用字表已写入 {out}({len(top)} 字)")


def build_parser():
    ap = argparse.ArgumentParser(
        prog="spec.py", description="规范查询 skill 唯一程序(两态: 查询态/维护态)")
    sub = ap.add_subparsers(dest="cmd", required=True)
    ap.add_argument("--config", help="config.json 路径(默认 skill 目录下)")

    p = sub.add_parser("index", help="加书/建索引: 质量检测→(OCR)→目录→切章→条文索引")
    p.add_argument("pdfs", nargs="*",
                   help="源文件路径(多个,支持 pdf/doc/docx/xls/xlsx/wps/et/ofd/png/tif/txt)")
    p.add_argument("--all", action="store_true",
                   help="库目录全部支持的文件(增量,已索引且未变则跳过)")
    p.add_argument("--only", nargs="+", metavar="EXT",
                   help="只处理指定扩展名(如 --only docx png,与 --all 或显式路径联用)")
    p.add_argument("--force", action="store_true", help="强制重建")
    p.add_argument("--rebuild", action="store_true",
                   help="从已有文本(ocr/extracted)重建目录/章/条文,不重 OCR(配合 set-page 视觉修正后生效)")
    p.add_argument("--jobs", type=int, default=1,
                   help="并行度(默认 1;COM 转换格式 doc/xls/wps/et 固定串行)")
    p.set_defaults(func=cmd_index)

    p = sub.add_parser("list", help="读书架")
    p.add_argument("-t", "--type", choices=["text", "ocr"], help="按类型过滤")
    p.add_argument("-c", "--category", help="按分类过滤")
    p.add_argument("-q", "--query", help="按 id/书名/规范号/别名过滤")
    p.set_defaults(func=cmd_list)

    p = sub.add_parser("toc", help="打印章节索引表")
    p.add_argument("book")
    p.add_argument("--chapter", type=int, help="只看第 N 章一行")
    p.set_defaults(func=cmd_toc)

    p = sub.add_parser("clause", help="条文号直查(未命中给出相邻条文)")
    p.add_argument("book")
    p.add_argument("clause")
    p.add_argument("--near", type=int, default=3, help="相邻条文数")
    p.set_defaults(func=cmd_clause)

    p = sub.add_parser("read", help="按页读原文(兜底;text 书自动补缓存)")
    p.add_argument("book")
    p.add_argument("page", type=int)
    p.add_argument("end", nargs="?", type=int, help="末页(单次最多 20 页)")
    p.set_defaults(func=cmd_read)

    p = sub.add_parser("grep", help="在章节文件上渐进式检索")
    p.add_argument("book", nargs="?", help="书 id;与 --all 二选一")
    p.add_argument("pattern", help="正则")
    p.add_argument("-i", "--ignore-case", action="store_true")
    p.add_argument("--ctx", type=int, default=1, help="上下文行数")
    p.add_argument("--max", type=int, default=30, help="命中上限(默认 30;0 = 不限,用于列命中书清单)")
    p.add_argument("--all", dest="all_books", action="store_true", help="跨全部已索引书")
    p.set_defaults(func=cmd_grep)

    p = sub.add_parser("ocr", help="整本/页范围批量 OCR(维护态)")
    p.add_argument("book")
    p.add_argument("--start", type=int)
    p.add_argument("--end", type=int)
    p.add_argument("--force", action="store_true", help="重跑已 OCR 页")
    p.set_defaults(func=cmd_ocr)

    p = sub.add_parser("img", help="渲染源 PDF 页为 PNG,供 agent 视觉复核/看图")
    p.add_argument("book")
    p.add_argument("page", type=int)
    p.add_argument("end", nargs="?", type=int, help="末页(缺省 = page;单页)")
    p.add_argument("--dpi", type=int, default=300)
    p.set_defaults(func=cmd_img)

    p = sub.add_parser("recheck", help="报告低置信/失败页(视觉复核清单)")
    p.add_argument("book")
    p.add_argument("--max", type=int, default=20, help="低置信条文最多列几条")
    p.set_defaults(func=cmd_recheck)

    p = sub.add_parser("set-page", help="回写某页被修正文本;配合 `index --rebuild` 生效")
    p.add_argument("book")
    p.add_argument("page", type=int)
    p.add_argument("--text", default=None, help="修正文本(缺省读 stdin)")
    p.set_defaults(func=cmd_set_page)

    p = sub.add_parser("status", help="库一致性检查 + 书架健康")
    p.set_defaults(func=cmd_status)

    p = sub.add_parser("remove", help="删除索引/登记,或标记被替代")
    p.add_argument("book")
    p.add_argument("--mark-superseded", metavar="NEW_ID", help="标记被新书替代(不删除)")
    p.set_defaults(func=cmd_remove)

    p = sub.add_parser("update-chars", help="生成/刷新常用字表(乱码检测资源)")
    p.add_argument("--from-pdfs", nargs="+", help="来源 PDF(干净文字版)")
    p.set_defaults(func=cmd_update_chars)
    return ap


def main():
    args = build_parser().parse_args()
    cfg = load_config(args)
    cfg["force"] = getattr(args, "force", False)
    cfg["rebuild"] = getattr(args, "rebuild", False)
    if args.cmd == "grep" and args.max is None:
        args.max = int(cfg.get("grep_default_max", 30))
    args.func(args, cfg)


if __name__ == "__main__":
    main()
